import pdfplumber
from docx import Document
import docx
import pytesseract
from PIL import Image
import os

# Extract text from PDF
def extract_text_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    print(text.strip())
    return text.strip()

# Extract text from DOCX
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_text))  # Format table rows with a separator

    return "\n".join(full_text).strip()

# Extract text from image (OCR)
def extract_text_from_image(file_path):
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)
    return text.strip()

# Detect file type and extract text
def extract_text(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".png", ".jpg", ".jpeg"]:
        return extract_text_from_image(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    else:
        return None


print(extract_text_from_pdf("./YB.pdf"))
#print(extract_text_from_docx("./fullstack.docx"))
#print(extract_text_from_image("./YB.png"))